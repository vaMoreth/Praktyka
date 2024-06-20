import os
from datetime import datetime
from flask import Flask, request, render_template, send_file
from docx import Document
from docx2pdf import convert

app = Flask(__name__)

months = {
    1: "січня", 2: "лютого", 3: "березня", 4: "квітня",
    5: "травня", 6: "червня", 7: "липня", 8: "серпня",
    9: "вересня", 10: "жовтня", 11: "листопада", 12: "грудня"
}

def romanize(num):
    num_map = [(5, 'V'), (4, 'IV'), (1, 'I')]
    roman = ''
    while num > 0:
        for i, r in num_map:
            while num >= i:
                roman += r
                num -= i
    return roman

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/generate', methods=['POST'])
def generate():
    doc_number = request.form['doc_number']
    issue_date = request.form['issue_date']
    course = romanize(int(request.form['course']))
    name = request.form['name']
    days = request.form['days']
    start_date = request.form['start_date']
    end_date = request.form['end_date']
    format_type = request.form['format']

    template_path = 'static/InputFiles/file2.docx'
    document = Document(template_path)

    replacements = {
        '{{doc_number}}': doc_number,
        '{{issue_date}}': format_date(issue_date, 'issue'),
        '{{course}}': course,
        '{{name}}': name,
        '{{days}}': days,
        '{{start_date}}': format_date(start_date, 'start'),
        '{{end_date}}': format_date(end_date, 'end')
    }

    process_document(document, replacements)


    if format_type == 'docx':
        output_path = os.path.join('uploads', f'Довідка_{name}_{doc_number}.docx')
        document.save(output_path)
    elif format_type == 'pdf':
        docx_path = os.path.join('uploads', f'Довідка_{name}_{doc_number}.docx')
        pdf_path = os.path.join('uploads', f'Довідка_{name}_{doc_number}.pdf')
        document.save(docx_path)
        convert(docx_path, pdf_path)
        output_path = pdf_path

    return send_file(output_path, as_attachment=True)


def format_date(date_str, date_type):
    date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    if date_type == 'start':
        return f'від «{date_obj.day}» {months[date_obj.month]} {date_obj.year} року'
    elif date_type == 'end':
        return f'по «{date_obj.day}» {months[date_obj.month]} {date_obj.year} року'
    elif date_type == 'issue':
        return f'«{date_obj.day}» {months[date_obj.month]} {date_obj.year} року'


def process_document(document, replacements):
    for paragraph in document.paragraphs:
        replace_placeholders(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders(paragraph, replacements)


def replace_placeholders(paragraph, replacements):
    for key, val in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, val)


if __name__ == '__main__':
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    app.run(debug=True)
